"""报表样式精化单测：文档信息、覆盖范围表、页眉。"""
from __future__ import annotations

import pytest

from weldai.domain.enums import ProcedureType
from weldai.services import export_procedure_docx


def test_docx_contains_manufacturer(tmp_path):
    """Word 应包含编制单位字段。"""
    from weldai.tests_demo import build_demo_pqr
    from docx import Document

    pqr = build_demo_pqr()
    pqr.manufacturer = "测试容器制造有限公司"
    pqr.project_no = "RCP-2026-001"
    out = export_procedure_docx(pqr, tmp_path / "mfg.docx")
    doc = Document(str(out))
    text = "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    assert "测试容器制造有限公司" in text
    assert "RCP-2026-001" in text


def test_docx_pqr_has_coverage_table(tmp_path):
    """PQR 应包含评定覆盖范围表。"""
    from weldai.tests_demo import build_demo_pqr
    from docx import Document

    pqr = build_demo_pqr()
    out = export_procedure_docx(pqr, tmp_path / "cov.docx")
    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    text += "\n" + "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    assert "覆盖范围" in text
    assert "NB/T 47014" in text


def test_docx_wps_no_coverage_table(tmp_path):
    """WPS 不应包含覆盖范围表（覆盖范围属 PQR）。"""
    from weldai.tests_demo import build_demo_wps
    from docx import Document

    wps = build_demo_wps()
    out = export_procedure_docx(wps, tmp_path / "wps.docx")
    doc = Document(str(out))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "评定覆盖范围" not in text


def test_docx_signature_has_signers(tmp_path):
    """签字栏应反映编制/审核/批准人。"""
    from weldai.tests_demo import build_demo_pqr
    from docx import Document

    pqr = build_demo_pqr()
    pqr.prepared_by = "张工"
    pqr.approved_by = "李总工"
    pqr.prepare_date = "2026-06-24"
    out = export_procedure_docx(pqr, tmp_path / "sig.docx")
    doc = Document(str(out))
    text = "\n".join(c.text for t in doc.tables for r in t.rows for c in r.cells)
    assert "张工" in text
    assert "李总工" in text
    assert "2026-06-24" in text


def test_docx_header_has_doc_no(tmp_path):
    """页眉应包含文档编号（多页追溯）。"""
    from weldai.tests_demo import build_demo_pqr

    pqr = build_demo_pqr()
    out = export_procedure_docx(pqr, tmp_path / "hdr.docx")
    # 页眉存储在 section.header
    import zipfile
    with zipfile.ZipFile(out) as z:
        header_files = [n for n in z.namelist() if "header" in n and n.endswith(".xml")]
        assert len(header_files) >= 1
        content = z.read(header_files[0]).decode("utf-8")
        assert "PQR-001" in content


def test_procedure_roundtrip_with_doc_meta(tmp_path):
    """文档信息字段应能通过仓储往返持久化。"""
    from weldai.persistence import (
        ProcedureRepository, get_session, init_db,
    )
    from weldai.standards import get_default_standard
    from weldai.tests_demo import build_demo_pqr

    init_db(":memory:")
    repo = ProcedureRepository(get_session(), get_default_standard())
    pqr = build_demo_pqr()
    pqr.manufacturer = "往返测试公司"
    pqr.prepared_by = "王工"
    repo.save(pqr)

    loaded = repo.get("PQR-001")
    assert loaded.manufacturer == "往返测试公司"
    assert loaded.prepared_by == "王工"
