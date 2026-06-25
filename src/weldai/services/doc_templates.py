"""标准 WPS/PQR 文档模板（NB/T 47015 一体化表格格式）。

输出 Word(.docx)：python-docx，单张合并单元格的一体化表格，
布局对标 NB/T 47015《压力容器焊接规程》正式表格：
  - 标题行（类型 + 文件编号，贯通）
  - 单位/方法/位置/热处理（双栏）
  - 母材（左）↔ 焊接位置/焊后热处理（右）
  - 焊材（左）↔ 坡口详图（右，嵌入矢量图）
  - 焊接参数表（贯通全宽，含每道焊接方法——支持组合焊）
  - 签字栏（编制/审核/批准）

坡口图通过 groove_renderer 按真实参数绘制后嵌入。
"""
from __future__ import annotations

import tempfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from ..domain.enums import ProcedureType
from ..domain.procedure import Procedure
from .groove_renderer import render_groove


# ---------------------------------------------------------------------------
# 公共：数值格式化与文档类型标题
# ---------------------------------------------------------------------------

def _type_title(t: ProcedureType) -> str:
    return {
        ProcedureType.WPS: "焊 接 工 艺 规 程 (WPS)",
        ProcedureType.PQR: "焊接工艺评定记录 (PQR)",
        ProcedureType.PWPS: "预焊接工艺规程 (pWPS)",
    }.get(t, "焊接工艺文件")


def _fmt(v, unit: str = "") -> str:
    if v is None or v == "":
        return ""
    if isinstance(v, float) and v.is_integer():
        v = int(v)
    return f"{v}{unit}"


def _groove_text(p: Procedure) -> str:
    if not p.joints:
        return "（无接头数据）"
    j = p.joints[0]
    g = j.groove
    return (f"{j.type.cn} | 坡口{g.type}形 | 角度{_fmt(g.angle, '°')} | "
            f"钝边{_fmt(g.root_face, 'mm')} | 间隙{_fmt(g.root_gap, 'mm')}"
            + (" | 带衬垫" if g.has_backing else ""))


# ---------------------------------------------------------------------------
# Word (.docx) 导出 —— 一体化表格
# ---------------------------------------------------------------------------

# 一体化表格列数（与布局对应）
_NCOLS = 6


def export_procedure_docx(procedure: Procedure, output_path: str | Path) -> Path:
    """导出工艺文件为 Word(.docx)，NB/T 47015 一体化表格格式，含坡口图。"""
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()

    _set_default_font(doc)
    _set_page(doc, procedure.doc_no)

    # 标题
    title = doc.add_heading(_type_title(procedure.type), level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _build_unified_table(doc, procedure)

    # 评定覆盖范围（仅 PQR）
    if procedure.type == ProcedureType.PQR:
        _add_coverage_table(doc, procedure)

    doc.save(str(out))
    return out


def _build_unified_table(doc: Document, p: Procedure) -> None:
    """构造一体化大表格。先按最大行数建空表，再逐格填内容并合并。"""
    # 估算行数：标题1 + 表头3 + 母材(1+max) + 预热1 + 焊材(1+max) + 坡口1 + 参数(1+max) + 签字1
    n_bm = max(len(p.base_metals), 1)
    n_cons = max(len(p.consumables), 1)
    n_pass = max(len(p.passes), 1)
    total_rows = 1 + 3 + (1 + n_bm) + 1 + (1 + n_cons) + 1 + (1 + n_pass) + 1
    table = doc.add_table(rows=total_rows, cols=_NCOLS)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    r = 0
    # ---- 行0：文件编号（贯通）----
    _merge_row(table, r, 0, _NCOLS, f"文件编号：{p.doc_no}", bold=True, bg="1565C0",
               color="FFFFFF")
    r += 1
    # ---- 行1-3：表头双栏 ----
    _set_cell(table, r, 0, "编制单位", bold=True)
    _set_cell(table, r, 1, p.manufacturer or "")
    _set_cell(table, r, 2, "项目编号", bold=True)
    _set_cell(table, r, 3, p.project_no or "")
    _set_cell(table, r, 4, "编制日期", bold=True)
    _set_cell(table, r, 5, p.prepare_date or "")
    r += 1
    _set_cell(table, r, 0, "焊接方法", bold=True)
    _set_cell(table, r, 1, p.process_display)  # 组合焊列出全部方法
    _set_cell(table, r, 2, "机械化程度", bold=True)
    _set_cell(table, r, 3, p.mechanization.cn)
    _set_cell(table, r, 4, "依据PQR", bold=True)
    _set_cell(table, r, 5, p.supporting_pqr_no or "—")
    r += 1
    _set_cell(table, r, 0, "评定标准", bold=True)
    _set_cell(table, r, 1, p.standard_version or "—")
    _set_cell(table, r, 2, "产品图号", bold=True)
    _set_cell(table, r, 3, p.drawing_no or "")
    _set_cell(table, r, 4, "冲击试验", bold=True)
    _set_cell(table, r, 5, "是" if p.impact_required else "否")
    r += 1

    # ---- 母材区（左3列）+ 焊接位置/预热（右3列）----
    bm_header_row = r
    _set_cell(table, r, 0, "母材", bold=True)
    _set_cell(table, r, 1, "类别号/组别号", bold=True)
    _set_cell(table, r, 2, "牌号/厚度", bold=True)
    _set_cell(table, r, 3, "焊接位置", bold=True)
    pos_text = ", ".join(pos.value for pos in p.positions) or "—"
    _merge_cells(table, r, 3, r, 5)  # 右3列合并显示位置
    _set_cell(table, r, 3, pos_text)
    r += 1
    for bm in p.base_metals:
        m = bm.metal
        _set_cell(table, r, 0, "")
        _set_cell(table, r, 1, f"{m.category}\n{m.group.group}")
        _set_cell(table, r, 2, f"{m.grade} / {_fmt(bm.thickness, 'mm')}")
        r += 1
    if not p.base_metals:
        _set_cell(table, r, 0, "（无）")
        r += 1

    # 预热/厚度范围/焊后热处理
    pwht = p.pwht
    pwht_desc = "无"
    if pwht.applied:
        pwht_desc = (f"{pwht.pwht_type} {_fmt(pwht.temp_min)}~{_fmt(pwht.temp_max, '℃')}"
                     f" 保温{_fmt(pwht.hold_time, 'min')}")
    _set_cell(table, r, 0, "母材厚度范围", bold=True)
    _set_cell(table, r, 1, _fmt(p.deposited_thickness, "mm"))
    _set_cell(table, r, 2, "最低预热", bold=True)
    _set_cell(table, r, 3, _fmt(p.preheat_min, "℃") or "无")
    _set_cell(table, r, 4, "焊后热处理", bold=True)
    _set_cell(table, r, 5, pwht_desc)
    r += 1

    # ---- 焊材区（左3列）+ 坡口（右3列，跨行）----
    cons_header_row = r
    _set_cell(table, r, 0, "焊接材料", bold=True)
    _set_cell(table, r, 1, "型号/分类栏位", bold=True)
    _set_cell(table, r, 2, "牌号/直径", bold=True)
    _set_cell(table, r, 3, "坡口形式与尺寸", bold=True)
    r += 1
    for c in p.consumables:
        _set_cell(table, r, 0, "")
        _set_cell(table, r, 1, f"{c.model}\n{c.classification_slot}")
        _set_cell(table, r, 2, f"{c.brand} / {_fmt(c.diameter, 'mm')}")
        r += 1
    if not p.consumables:
        _set_cell(table, r, 0, "（无）")
        r += 1

    # 坡口图行（右3列合并嵌入图，左3列写文字说明）
    groove_row = r
    _set_cell(table, r, 0, "坡口详图", bold=True)
    _merge_cells(table, r, 1, r, 2)
    _set_cell(table, r, 1, _groove_text(p))
    _merge_cells(table, r, 3, r, 5)
    _embed_groove_image(table, r, 3, p)
    r += 1

    # ---- 焊接参数表（贯通全宽，含每道焊接方法——支持组合焊）----
    param_header_row = r
    _set_cell(table, r, 0, "焊道", bold=True)
    _set_cell(table, r, 1, "焊接方法", bold=True)
    _merge_cells(table, r, 1, r, 1)
    _set_cell(table, r, 2, "层次", bold=True)
    _set_cell(table, r, 3, "焊材/直径", bold=True)
    _set_cell(table, r, 4, "电流(A)/极性", bold=True)
    _set_cell(table, r, 5, "电压(V)/速度", bold=True)
    r += 1
    for pa in p.passes:
        p_eff = pa.effective_process(p.process)
        cons = pa.consumable
        cons_disp = cons.brand if cons else ""
        if pa.diameter:
            cons_disp += f" /φ{pa.diameter:g}"
        cur = (f"{_fmt(pa.current_min)}~{_fmt(pa.current_max)}"
               if (pa.current_min or pa.current_max) else "")
        pol = pa.current_type.cn if pa.current_type else ""
        vol = (f"{_fmt(pa.voltage_min)}~{_fmt(pa.voltage_max)}"
               if (pa.voltage_min or pa.voltage_max) else "")
        spd = _fmt(pa.travel_speed)
        _set_cell(table, r, 0, str(pa.sequence))
        _set_cell(table, r, 1, f"{p_eff.value} {p_eff.cn}")
        _set_cell(table, r, 2, pa.layer_role)
        _set_cell(table, r, 3, cons_disp)
        _set_cell(table, r, 4, f"{cur}\n{pol}")
        _set_cell(table, r, 5, f"{vol}\n{spd}cm/min")
        r += 1
    if not p.passes:
        _set_cell(table, r, 0, "（无焊道参数）")
        r += 1

    # ---- 签字栏（贯通全宽）----
    _set_cell(table, r, 0, "编制", bold=True)
    _set_cell(table, r, 1, p.prepared_by or "")
    _set_cell(table, r, 2, "审核", bold=True)
    _set_cell(table, r, 3, p.reviewed_by or "")
    _set_cell(table, r, 4, "批准", bold=True)
    _set_cell(table, r, 5, p.approved_by or "")

    # 标题行/参数表头/签字栏底纹
    _shade_row(table, 0, "1565C0", text_color="FFFFFF")
    _shade_row(table, param_header_row, "E3F2FD")


def _embed_groove_image(table, row: int, col: int, p: Procedure) -> None:
    """在指定单元格嵌入坡口矢量图。"""
    if not p.joints:
        _set_cell(table, row, col, "（无接头数据）")
        return
    joint = p.joints[0]
    thickness = joint.thickness or (p.base_metals[0].thickness if p.base_metals else 16)
    tmp = Path(tempfile.mktemp(suffix=".png"))
    try:
        render_groove(joint.groove, thickness, tmp)
        cell = table.cell(row, col)
        # 清空默认段落
        cell.text = ""
        run = cell.paragraphs[0].add_run()
        run.add_picture(str(tmp), width=Cm(6))
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception:
        _set_cell(table, row, col, _groove_text(p))
    finally:
        if tmp.exists():
            try:
                tmp.unlink()
            except Exception:
                pass


def _set_default_font(doc: Document) -> None:
    """设置默认中文字体为宋体。"""
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10.5)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:eastAsia"), "宋体")


def _set_page(doc: Document, doc_no: str = "") -> None:
    """A4 页面，标准页边距 + 页眉（多页追溯）。"""
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = section.right_margin = Cm(1.8)
    section.top_margin = section.bottom_margin = Cm(2)
    if doc_no:
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.text = f"文件编号：{doc_no}"
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in hp.runs:
            run.font.size = Pt(8)


def _set_cell(table, row: int, col: int, text, bold: bool = False,
              size: float = 9) -> None:
    """设置单元格文本（支持多行用\n）。"""
    cell = table.cell(row, col)
    cell.text = ""
    lines = str(text).split("\n") if text is not None else [""]
    for i, line in enumerate(lines):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run = para.add_run(line)
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = "Times New Roman"
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), "宋体")


def _merge_row(table, row: int, start_col: int, n_cols: int,
               text: str, bold: bool = False, bg: str = "",
               color: str = "") -> None:
    """合并整行单元格并设置文本。"""
    _merge_cells(table, row, start_col, row, start_col + n_cols - 1)
    cell = table.cell(row, start_col)
    _set_cell(table, row, start_col, text, bold=bold)
    if bg:
        _shade_cell(cell, bg)
    if color:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor.from_string(color)


def _merge_cells(table, r1: int, c1: int, r2: int, c2: int) -> None:
    """合并矩形区域单元格（r1,c1)-(r2,c2)。"""
    a = table.cell(r1, c1)
    b = table.cell(r2, c2)
    a.merge(b)


def _shade_cell(cell, hex_color: str) -> None:
    """给单元格设置底纹颜色。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        from docx.oxml import OxmlElement
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)


def _shade_row(table, row: int, hex_color: str,
               text_color: str = "") -> None:
    """给整行设置底纹。"""
    for c in range(_NCOLS):
        cell = table.cell(row, c)
        _shade_cell(cell, hex_color)
        if text_color:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor.from_string(text_color)


def _add_coverage_table(doc: Document, p: Procedure) -> None:
    """评定覆盖范围表（NB/T 47015 要求的厚度/管径/位置覆盖范围栏）。

    仅 PQR 类型输出此表，展示本评定可覆盖的生产范围。
    """
    from ..engine import CoverageEngine
    from ..standards import get_default_standard

    doc.add_heading("评定覆盖范围（本PQR可覆盖的生产参数范围）", level=2)
    std = get_default_standard()
    cov = CoverageEngine(std)
    trange = cov.thickness_coverage(p)
    prange = cov.diameter_coverage(p)
    positions = [pos.value for pos in cov.position_coverage(p)]

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"

    def _fmt_range(lo, hi):
        hi_str = f"{hi:g}" if hi is not None else "不限"
        return f"{lo:g} ~ {hi_str} mm"

    rows = [
        ("母材厚度覆盖范围",
         _fmt_range(trange.min_t, trange.max_t) if trange else "—"),
        ("管径覆盖范围",
         _fmt_range(prange.min_d, prange.max_d) if prange else "（板材/无管径要求）"),
        ("焊接位置覆盖", ", ".join(positions) if positions else "—"),
        ("依据标准", f"{std.standard_code} 表7"),
    ]
    for i, (label_text, val) in enumerate(rows):
        _set_cell(table, i, 0, label_text, bold=True)
        _set_cell(table, i, 1, val)
    note = doc.add_paragraph(
        "注：覆盖范围依据 NB/T 47014 表7，冲击试验要求会影响厚度上限。"
        "生产前请用 weldAI 规则引擎对具体 WPS 进行校验。"
    )
    for run in note.runs:
        run.font.size = Pt(9)


# ---------------------------------------------------------------------------
# 统一入口
# ---------------------------------------------------------------------------

def export_procedure(
    procedure: Procedure,
    output_path: str | Path,
    fmt: str = "docx",
) -> Path:
    """统一导出入口。

    fmt: 'docx'（Word，可编辑）或 'pdf'（归档）。
    """
    fmt = fmt.lower()
    if fmt == "docx":
        return export_procedure_docx(procedure, output_path)
    if fmt == "pdf":
        from .export_service import export_procedure_pdf
        return export_procedure_pdf(procedure, output_path)
    raise ValueError(f"不支持的格式: {fmt}（支持 docx / pdf）")
