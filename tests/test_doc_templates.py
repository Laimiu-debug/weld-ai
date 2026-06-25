"""标准文档模板单测。"""
from __future__ import annotations

import pytest

from weldai.services import export_procedure, export_procedure_docx


def test_export_docx_generates_valid_file(tmp_path):
    """Word 导出应生成有效 .docx 文件。"""
    from weldai.tests_demo import build_demo_pqr

    out = export_procedure_docx(build_demo_pqr(), tmp_path / "test.docx")
    assert out.exists()
    assert out.stat().st_size > 5000  # 含图片，应较大
    # .docx 是 zip 格式，文件头应为 PK
    with open(out, "rb") as f:
        assert f.read(2) == b"PK"


def test_export_docx_contains_key_content(tmp_path):
    """Word 内容应包含关键字段（编号/母材/焊材）。"""
    from weldai.tests_demo import build_demo_pqr
    from docx import Document

    pqr = build_demo_pqr()
    out = export_procedure_docx(pqr, tmp_path / "check.docx")
    doc = Document(str(out))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text += "\n" + cell.text
    assert "PQR-001" in full_text
    assert "Q345R" in full_text
    assert "J507" in full_text


def test_export_docx_has_tables(tmp_path):
    """Word 应包含一体化主表格（NB/T 47015 单表格式）+ PQR覆盖表。"""
    from weldai.tests_demo import build_demo_pqr
    from docx import Document

    # PQR 会额外输出覆盖范围表（主表 + 覆盖表 = 2）
    out = export_procedure_docx(build_demo_pqr(), tmp_path / "tables.docx")
    doc = Document(str(out))
    # 一体化主表格至少存在
    assert len(doc.tables) >= 1
    # 主表格行数应足够多（含表头/母材/焊材/焊道/签字多区段）
    main_table = doc.tables[0]
    assert len(main_table.rows) >= 8


def test_export_docx_has_groove_image(tmp_path):
    """Word 应嵌入坡口图片（至少一张图）。"""
    from weldai.tests_demo import build_demo_pqr
    from docx import Document

    out = export_procedure_docx(build_demo_pqr(), tmp_path / "img.docx")
    doc = Document(str(out))
    # 统计内嵌图片
    import zipfile
    with zipfile.ZipFile(out) as z:
        images = [n for n in z.namelist() if n.startswith("word/media/")]
    assert len(images) >= 1


def test_export_unified_entry_docx(tmp_path):
    """统一入口 export_procedure(fmt='docx') 应等同 docx 导出。"""
    from weldai.tests_demo import build_demo_pqr

    out = export_procedure(build_demo_pqr(), tmp_path / "uni.docx", fmt="docx")
    assert out.exists()


def test_export_unified_entry_pdf(tmp_path):
    """统一入口 fmt='pdf' 应生成 PDF。"""
    from weldai.tests_demo import build_demo_pqr

    out = export_procedure(build_demo_pqr(), tmp_path / "uni.pdf", fmt="pdf")
    assert out.exists()
    with open(out, "rb") as f:
        assert f.read(4) == b"%PDF"


def test_export_invalid_format_raises(tmp_path):
    """不支持的格式应抛错。"""
    from weldai.tests_demo import build_demo_pqr

    with pytest.raises(ValueError):
        export_procedure(build_demo_pqr(), tmp_path / "x.txt", fmt="txt")


def test_export_docx_wps_type(tmp_path):
    """WPS 类型应有对应标题。"""
    from weldai.tests_demo import build_demo_wps
    from docx import Document

    out = export_procedure_docx(build_demo_wps(), tmp_path / "wps.docx")
    doc = Document(str(out))
    title_text = doc.paragraphs[0].text
    assert "WPS" in title_text
